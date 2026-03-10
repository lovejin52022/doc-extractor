"""Document Parser Service - 文档解析服务"""
import io
from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Optional, Dict, Any, Union
import tempfile


class DocumentParser(ABC):
    """文档解析器抽象基类"""
    
    @abstractmethod
    def parse(self, file_path: str) -> str:
        """解析文档，返回文本内容"""
        pass
    
    @abstractmethod
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取文档元数据"""
        pass


class PDFParser(DocumentParser):
    """PDF 文档解析器"""
    
    def __init__(self):
        self._client = None
    
    def _get_client(self):
        if self._client is None:
            try:
                import PyPDF2
                self._client = PyPDF2
            except ImportError:
                raise ImportError("PyPDF2 库未安装，请运行: pip install PyPDF2")
        return self._client
    
    def parse(self, file_path: str) -> str:
        """解析 PDF 文件"""
        client = self._get_client()
        
        text_parts = []
        
        with open(file_path, 'rb') as f:
            reader = client.PdfReader(f)
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text:
                    text_parts.append(f"[Page {page_num}]\n{text}")
        
        return "\n\n".join(text_parts)
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取 PDF 元数据"""
        client = self._get_client()
        
        metadata = {
            "file_name": Path(file_path).name,
            "file_size": Path(file_path).stat().st_size,
            "pages": 0
        }
        
        try:
            with open(file_path, 'rb') as f:
                reader = client.PdfReader(f)
                metadata["pages"] = len(reader.pages)
                
                # 提取 PDF 元信息
                if reader.metadata:
                    pdf_meta = reader.metadata
                    metadata["title"] = pdf_meta.get("/Title")
                    metadata["author"] = pdf_meta.get("/Author")
                    metadata["subject"] = pdf_meta.get("/Subject")
                    metadata["creator"] = pdf_meta.get("/Creator")
                    metadata["producer"] = pdf_meta.get("/Producer")
                    metadata["creation_date"] = str(pdf_meta.get("/CreationDate", ""))
                    metadata["modification_date"] = str(pdf_meta.get("/ModDate", ""))
        except Exception as e:
            metadata["error"] = str(e)
        
        return metadata


class DOCXParser(DocumentParser):
    """DOCX 文档解析器"""
    
    def __init__(self):
        self._client = None
    
    def _get_client(self):
        if self._client is None:
            try:
                import docx
                self._client = docx
            except ImportError:
                raise ImportError("python-docx 库未安装，请运行: pip install python-docx")
        return self._client
    
    def parse(self, file_path: str) -> str:
        """解析 DOCX 文件"""
        client = self._get_client()
        
        text_parts = []
        
        doc = client.Document(file_path)
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取 DOCX 元数据"""
        client = self._get_client()
        
        metadata = {
            "file_name": Path(file_path).name,
            "file_size": Path(file_path).stat().st_size,
            "paragraphs": 0,
            "tables": 0
        }
        
        try:
            doc = client.Document(file_path)
            metadata["paragraphs"] = len(doc.paragraphs)
            metadata["tables"] = len(doc.tables)
            
            # 提取核心属性
            core_props = doc.core_properties
            metadata["title"] = core_props.title
            metadata["author"] = core_props.author
            metadata["subject"] = core_props.subject
            metadata["keywords"] = core_props.keywords
            metadata["created"] = str(core_props.created) if core_props.created else None
            metadata["modified"] = str(core_props.modified) if core_props.modified else None
        except Exception as e:
            metadata["error"] = str(e)
        
        return metadata


class TextParser(DocumentParser):
    """纯文本解析器"""
    
    def parse(self, file_path: str) -> str:
        """解析文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取文本文件元数据"""
        path = Path(file_path)
        
        metadata = {
            "file_name": path.name,
            "file_size": path.stat().st_size,
            "encoding": "utf-8"
        }
        
        # 尝试检测编码
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                metadata["lines"] = content.count('\n') + 1
                metadata["characters"] = len(content)
        except UnicodeDecodeError:
            metadata["encoding"] = "unknown"
        
        return metadata


class MarkdownParser(DocumentParser):
    """Markdown 解析器"""
    
    def parse(self, file_path: str) -> str:
        """解析 Markdown 文件"""
        # 简单实现：直接读取文本
        # 实际可以使用 markdown 库进行转换
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取 Markdown 元数据"""
        path = Path(file_path)
        
        metadata = {
            "file_name": path.name,
            "file_size": path.stat().st_size,
            "type": "markdown"
        }
        
        # 尝试提取 YAML 头
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
                if content.startswith('---'):
                    parts = content.split('---', 2)
                    if len(parts) >= 3:
                        frontmatter = parts[1]
                        metadata["frontmatter"] = frontmatter
                
                # 统计 Markdown 元素
                metadata["headers"] = content.count('# ')
                metadata["code_blocks"] = content.count('```')
                metadata["links"] = content.count('](')
        except Exception as e:
            metadata["error"] = str(e)
        
        return metadata


def create_parser(file_path: str) -> DocumentParser:
    """工厂函数：创建合适的文档解析器"""
    path = Path(file_path)
    extension = path.suffix.lower()
    
    parsers = {
        ".pdf": PDFParser,
        ".docx": DOCXParser,
        ".doc": DOCXParser,
        ".txt": TextParser,
        ".md": MarkdownParser,
        ".markdown": MarkdownParser,
    }
    
    parser_class = parsers.get(extension)
    
    if not parser_class:
        raise ValueError(f"不支持的文件类型: {extension}")
    
    return parser_class()


def parse_document(file_path: str) -> str:
    """解析文档的便捷函数"""
    parser = create_parser(file_path)
    return parser.parse(file_path)


def extract_document_metadata(file_path: str) -> Dict[str, Any]:
    """提取文档元数据的便捷函数"""
    parser = create_parser(file_path)
    return parser.extract_metadata(file_path)


class ChunkStrategy:
    """文档分块策略"""
    
    @staticmethod
    def fixed_size(
        text: str,
        chunk_size: int = 500,
        overlap: int = 50
    ) -> List[str]:
        """固定大小分块"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            
            # 尝试在句子边界分割
            if end < len(text):
                # 找最后一个句号或逗号
                last_period = max(chunk.rfind('。'), chunk.rfind('.'), chunk.rfind('?'))
                last_comma = max(chunk.rfind('，'), chunk.rfind(','))
                
                if last_period > chunk_size // 2:
                    end = start + last_period + 1
                elif last_comma > chunk_size // 2:
                    end = start + last_comma + 1
            
            chunks.append(text[start:end].strip())
            start = end - overlap
        
        return [c for c in chunks if c]
    
    @staticmethod
    def by_paragraph(text: str, min_chunk_size: int = 50) -> List[str]:
        """按段落分块"""
        # 按换行符分割
        paragraphs = text.split('\n')
        
        chunks = []
        current_chunk = []
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            current_chunk.append(para)
            
            # 如果当前块足够大，保存并重置
            if sum(len(p) for p in current_chunk) >= min_chunk_size:
                chunks.append('\n'.join(current_chunk))
                current_chunk = []
        
        # 添加最后一个块
        if current_chunk:
            chunks.append('\n'.join(current_chunk))
        
        return chunks
    
    @staticmethod
    def recursive(
        text: str,
        separators: List[str] = None,
        min_chunk_size: int = 50
    ) -> List[str]:
        """递归分块（先大块，后小块）"""
        if separators is None:
            separators = ['\n\n', '\n', '。', '.', '，', ',', ' ']
        
        def split_text(txt: str, sep_list: List[str]) -> List[str]:
            if not sep_list:
                return [txt] if len(txt) >= min_chunk_size else []
            
            sep = sep_list[0]
            parts = txt.split(sep)
            
            result = []
            current = []
            
            for part in parts:
                current.append(part)
                combined = sep.join(current)
                
                if len(combined) >= min_chunk_size:
                    result.append(combined)
                    current = []
            
            if current:
                # 递归处理剩余部分
                remaining = sep.join(current)
                if remaining:
                    result.extend(split_text(remaining, sep_list[1:]))
            
            return result
        
        return split_text(text, separators)
